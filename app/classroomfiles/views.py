from django.shortcuts import render
from rest_framework import viewsets,generics
from .models import ClassRoomFiles
from .serializers import ClassRoomFilesSerializer
from rest_framework import filters
from rest_framework import permissions
from rest_framework.response import Response
from rest_framework import status, viewsets
from users.models import User
from docx import Document
import docx2txt
from users.serializers import GetUserSerializer
import pusher
import io
from decouple import config
import re
import urllib.request
from nltk.util import ngrams, pad_sequence, everygrams
from nltk.tokenize import word_tokenize
from nltk.lm import MLE, WittenBellInterpolated
import nltk
import numpy as np
import plotly.graph_objects as go
from scipy.ndimage import gaussian_filter

nltk.download('punkt')
class ClassRoomFilesView(viewsets.ModelViewSet):
    filter_backends = [filters.SearchFilter]
    search_fields = ['location']
    queryset=ClassRoomFiles.objects.all()
    permissions_class = [permissions.AllowAny]
    serializer_class=ClassRoomFilesSerializer

    def create(self,request):
        res = request.data
        s = ClassRoomFilesSerializer(data=res)
        s.is_valid(raise_exception=True)
        s.save()
        if(res.get('category')=='Modules'):
            return Response()
        # print(open(res.get('files').path).read())
        listResults = []
        items = ClassRoomFiles.objects.filter(folder_id=res.get('folder_id'))
        items = ClassRoomFilesSerializer(items,many=True)
        for x in items.data:
            if(len(items.data)>1 and x['user_id']!=s.data.get('user_id')):
                print(x['files'])
                listResults.append({"result":recognition(s.data.get('files'),x['files']),"name":x['author']})
        
        if(len(listResults)==0):
            print("okayy")
            ClassRoomFiles.objects.filter(id=s.data['id']).update(results=0.0)
        else:
            listResult = []
            for x in listResults:
                listResult.append(x['result'])

            for x in listResults:
                if(max(listResult)==x['result']):
                    ClassRoomFiles.objects.filter(id=s.data['id']).update(results=max(listResult),percent_from=x['name'])
        return Response(data = {})
        


class GetFileByProfID(generics.GenericAPIView):
    def get(self,request,format=None,folder_id=None):
        items = ClassRoomFiles.objects.filter(folder_id=folder_id)
        items = ClassRoomFilesSerializer(items,many=True)
        return Response(data = items.data)




def recognition(url,testfiles):
    
    try:
        link = url
        f = urllib.request.urlopen(link) 
        s = urllib.request.urlopen(testfiles)   
        testFile = s.read()         
        myfile = f.read()
        train_val = ""
        test_val = ""
        train = Document(io.BytesIO(myfile))
        test = Document(io.BytesIO(testFile))
        for p in train.paragraphs:
            train_val = p.text
        for p in test.paragraphs:
            test_val = p.text
        
        if(train_val==''):
            train_val = docx2txt.process(io.BytesIO(myfile))
        if(test_val==''):
            test_val = docx2txt.process(io.BytesIO(testFile))


        print(train_val)
        train_text = re.sub(r"\[.*\]|\{.*\}", "", train_val)
        train_text = re.sub(r'[^\w\s]', "",train_val)
        n = 4
        training_data = list(pad_sequence(word_tokenize(train_text), n, 
                                        pad_left=True, 
                                        left_pad_symbol="<s>"))
        ngrams = list(everygrams(training_data, max_len=n))
        print(len(ngrams))
        model = WittenBellInterpolated(n)
        model.fit([ngrams], vocabulary_text=training_data)
        test_text = re.sub(r'[^\w\s]', "", test_val)
        testing_data = list(pad_sequence(word_tokenize(test_text), n, 
                                        pad_left=True,
                                        left_pad_symbol="<s>"))
        scores = []
        for i, item in enumerate(testing_data[n-1:]):
            s = model.score(item, testing_data[i:i+n-1])
            scores.append(s)
        scores_np = np.array(scores)
        print(scores_np)
        return max(scores_np)
    except Exception as e:
        print(e)
        return 0.0